from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'profile.jpg', 
    width = Inches(2.0)
)

# personal details
name = input('What is your name? ')
speak('Hello ' + name + ' how are you today?')
phone_number = input('What is your phone number? ')
email = input('What is your email? ')

document.add_paragraph(
    name + ' | ' + phone_number +' | ' + email)

# about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell about yourself. ')
)

# work experience
document.add_heading('Work experience ')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company+ ' ')
p.add_run(experience_details)

# more experince
while True:
    has_more_experinces = input(
        'do you have more experiences? Yes or No ')
    if has_more_experinces.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' - ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading('Skills')
skill = input('Enter skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('do you have more skills? ')
    if has_more_skills.lower() == 'yes':
        skill = input('skill? ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]

document.save('cv.docx')